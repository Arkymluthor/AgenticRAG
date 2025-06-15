import logging
import os
import tempfile
import json
import io
import re
from typing import Dict, Any, Union, List, Optional, Tuple
import azure.functions as func
from azure.storage.blob import BlobServiceClient, BlobClient
import requests
import pandas as pd
import numpy as np
from concurrent.futures import ThreadPoolExecutor, as_completed

from ..orchestrator import RetryableError

logger = logging.getLogger(__name__)

# Maximum file size to process in memory (100MB)
MAX_IN_MEMORY_SIZE = 100 * 1024 * 1024

# Maximum number of pages to process in a single batch
MAX_PAGES_PER_BATCH = 50

# Maximum number of concurrent extraction tasks
MAX_CONCURRENT_TASKS = 5


def run(input_data: Union[str, Dict[str, Any]]) -> str:
    """
    Extracts text from a document.
    
    Args:
        input_data: Either a blob URI string or a dictionary with blob_uri and file_type
        
    Returns:
        Extracted text content as a string
        
    Raises:
        RetryableError: For transient failures that should be retried
        ValueError: For invalid input or unsupported document types
    """
    # Parse input
    if isinstance(input_data, str):
        blob_uri = input_data
        file_type = None  # Will be detected
        page_range = None
    elif isinstance(input_data, dict):
        blob_uri = input_data.get("blob_uri")
        file_type = input_data.get("file_type")
        page_range = input_data.get("page_range")  # Optional page range for large files
    else:
        raise ValueError(f"Invalid input type: {type(input_data)}")
    
    if not blob_uri:
        raise ValueError("No blob URI provided")
    
    logger.info(f"Extracting text from {blob_uri} (type: {file_type or 'auto-detect'})")
    
    try:
        # Get blob size to determine processing approach
        blob_size = get_blob_size(blob_uri)
        logger.info(f"Blob size: {blob_size / (1024 * 1024):.2f} MB")
        
        # For large files, use streaming approach
        if blob_size > MAX_IN_MEMORY_SIZE and not page_range:
            logger.info(f"Large file detected, using streaming approach")
            return extract_large_file(blob_uri, file_type)
        
        # For smaller files or specific page ranges, download and process
        temp_file_path = download_blob(blob_uri)
        
        # Extract text based on file type
        if not file_type:
            # Detect file type if not provided
            file_extension = blob_uri.split('.')[-1].lower() if '.' in blob_uri else ''
            if file_extension in ['pdf', 'docx', 'doc']:
                file_type = 'pdf' if file_extension == 'pdf' else 'docx'
            elif file_extension in ['txt', 'text', 'md']:
                file_type = 'txt'
            else:
                file_type = 'scan'  # Default to scan for unknown types
        
        # Extract text based on file type
        if file_type == 'pdf':
            text, tables = extract_text_from_pdf(temp_file_path, page_range)
        elif file_type == 'docx':
            text, tables = extract_text_from_docx(temp_file_path)
        elif file_type == 'txt':
            text = extract_text_from_txt(temp_file_path)
            tables = []
        elif file_type == 'scan':
            text, tables = extract_text_from_scan(temp_file_path, blob_uri)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
        
        # Clean up temporary file
        try:
            os.remove(temp_file_path)
        except Exception as e:
            logger.warning(f"Failed to remove temporary file {temp_file_path}: {str(e)}")
        
        # Combine text and tables
        combined_text = combine_text_and_tables(text, tables)
        
        logger.info(f"Successfully extracted {len(combined_text)} characters from {blob_uri}")
        return combined_text
        
    except Exception as e:
        # Handle transient errors
        if is_transient_error(e):
            logger.warning(f"Transient error extracting text: {str(e)}")
            raise RetryableError(f"Transient error: {str(e)}")
        
        # Handle permanent errors
        logger.error(f"Error extracting text: {str(e)}")
        raise


def get_blob_size(blob_uri: str) -> int:
    """
    Get the size of a blob in bytes.
    
    Args:
        blob_uri: The URI of the blob
        
    Returns:
        Size of the blob in bytes
    """
    # Parse the blob URI
    parts = blob_uri.replace("https://", "").split("/")
    account_name = parts[0].split('.')[0]
    container_name = parts[1]
    blob_path = "/".join(parts[2:])
    
    # Get connection string from environment
    connection_string = os.environ.get("AzureWebJobsStorage")
    if not connection_string:
        raise ValueError("AzureWebJobsStorage connection string not found")
    
    # Create blob service client
    blob_service_client = BlobServiceClient.from_connection_string(connection_string)
    
    # Get blob client
    container_client = blob_service_client.get_container_client(container_name)
    blob_client = container_client.get_blob_client(blob_path)
    
    # Get blob properties
    properties = blob_client.get_blob_properties()
    
    return properties.size


def extract_large_file(blob_uri: str, file_type: Optional[str] = None) -> str:
    """
    Extract text from a large file using a streaming approach.
    
    Args:
        blob_uri: The URI of the blob
        file_type: Optional file type
        
    Returns:
        Extracted text
    """
    # Determine file type if not provided
    if not file_type:
        file_extension = blob_uri.split('.')[-1].lower() if '.' in blob_uri else ''
        if file_extension in ['pdf']:
            file_type = 'pdf'
        elif file_extension in ['docx', 'doc']:
            file_type = 'docx'
        elif file_extension in ['txt', 'text', 'md']:
            file_type = 'txt'
        else:
            file_type = 'scan'
    
    # For PDF files, use page-by-page processing
    if file_type == 'pdf':
        return extract_large_pdf(blob_uri)
    
    # For other file types, download and process normally
    temp_file_path = download_blob(blob_uri)
    
    try:
        if file_type == 'docx':
            text, tables = extract_text_from_docx(temp_file_path)
        elif file_type == 'txt':
            text = extract_text_from_txt(temp_file_path)
            tables = []
        elif file_type == 'scan':
            text, tables = extract_text_from_scan(temp_file_path, blob_uri)
        else:
            raise ValueError(f"Unsupported file type for large file extraction: {file_type}")
        
        # Combine text and tables
        combined_text = combine_text_and_tables(text, tables)
        return combined_text
    
    finally:
        # Clean up temporary file
        try:
            os.remove(temp_file_path)
        except Exception as e:
            logger.warning(f"Failed to remove temporary file {temp_file_path}: {str(e)}")


def extract_large_pdf(blob_uri: str) -> str:
    """
    Extract text from a large PDF file by processing it page by page.
    
    Args:
        blob_uri: The URI of the blob
        
    Returns:
        Extracted text
    """
    # Get page count
    page_count = get_pdf_page_count(blob_uri)
    logger.info(f"PDF has {page_count} pages")
    
    # Process in batches
    all_text = []
    all_tables = []
    
    # Create batches of pages
    batches = []
    for start_page in range(1, page_count + 1, MAX_PAGES_PER_BATCH):
        end_page = min(start_page + MAX_PAGES_PER_BATCH - 1, page_count)
        batches.append((start_page, end_page))
    
    # Process batches concurrently
    with ThreadPoolExecutor(max_workers=MAX_CONCURRENT_TASKS) as executor:
        # Submit tasks
        future_to_batch = {
            executor.submit(
                process_pdf_batch, 
                blob_uri, 
                batch[0], 
                batch[1]
            ): batch for batch in batches
        }
        
        # Process results as they complete
        for future in as_completed(future_to_batch):
            batch = future_to_batch[future]
            try:
                text, tables = future.result()
                all_text.append(text)
                all_tables.extend(tables)
                logger.info(f"Processed pages {batch[0]}-{batch[1]}")
            except Exception as e:
                logger.error(f"Error processing pages {batch[0]}-{batch[1]}: {str(e)}")
                # Continue with other batches
    
    # Combine all text and tables
    combined_text = combine_text_and_tables("\n\n".join(all_text), all_tables)
    
    return combined_text


def get_pdf_page_count(blob_uri: str) -> int:
    """
    Get the number of pages in a PDF file.
    
    Args:
        blob_uri: The URI of the blob
        
    Returns:
        Number of pages
    """
    # Download a small portion of the PDF to get page count
    temp_file_path = download_blob(blob_uri, full_download=False)
    
    try:
        import PyPDF2
        
        with open(temp_file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            return len(pdf_reader.pages)
    
    finally:
        # Clean up temporary file
        try:
            os.remove(temp_file_path)
        except Exception as e:
            logger.warning(f"Failed to remove temporary file {temp_file_path}: {str(e)}")


def process_pdf_batch(blob_uri: str, start_page: int, end_page: int) -> Tuple[str, List[Dict[str, Any]]]:
    """
    Process a batch of PDF pages.
    
    Args:
        blob_uri: The URI of the blob
        start_page: Starting page number (1-based)
        end_page: Ending page number (inclusive)
        
    Returns:
        Tuple of (extracted text, tables)
    """
    # Create input data with page range
    input_data = {
        "blob_uri": blob_uri,
        "file_type": "pdf",
        "page_range": (start_page, end_page)
    }
    
    # Download the blob
    temp_file_path = download_blob(blob_uri)
    
    try:
        # Extract text and tables
        text, tables = extract_text_from_pdf(temp_file_path, (start_page, end_page))
        return text, tables
    
    finally:
        # Clean up temporary file
        try:
            os.remove(temp_file_path)
        except Exception as e:
            logger.warning(f"Failed to remove temporary file {temp_file_path}: {str(e)}")


def download_blob(blob_uri: str, full_download: bool = True) -> str:
    """
    Download a blob to a temporary file.
    
    Args:
        blob_uri: The URI of the blob to download
        full_download: Whether to download the entire blob or just the first part
        
    Returns:
        Path to the temporary file
    """
    # Parse the blob URI
    parts = blob_uri.replace("https://", "").split("/")
    container_name = parts[1]
    blob_path = "/".join(parts[2:])
    
    # Get connection string from environment
    connection_string = os.environ.get("AzureWebJobsStorage")
    if not connection_string:
        raise ValueError("AzureWebJobsStorage connection string not found")
    
    # Create blob service client
    blob_service_client = BlobServiceClient.from_connection_string(connection_string)
    
    # Get blob client
    container_client = blob_service_client.get_container_client(container_name)
    blob_client = container_client.get_blob_client(blob_path)
    
    # Create a temporary file
    fd, temp_file_path = tempfile.mkstemp(suffix=f".{blob_path.split('.')[-1]}" if '.' in blob_path else "")
    os.close(fd)
    
    # Download the blob to the temporary file
    with open(temp_file_path, "wb") as file:
        if full_download:
            # Download the entire blob
            blob_data = blob_client.download_blob()
            file.write(blob_data.readall())
        else:
            # Download just the first part (for PDF header)
            blob_data = blob_client.download_blob(offset=0, length=10240)
            file.write(blob_data.readall())
    
    return temp_file_path


def extract_text_from_pdf(file_path: str, page_range: Optional[Tuple[int, int]] = None) -> Tuple[str, List[Dict[str, Any]]]:
    """
    Extract text and tables from a PDF file.
    
    Args:
        file_path: Path to the PDF file
        page_range: Optional tuple of (start_page, end_page) to extract (1-based)
        
    Returns:
        Tuple of (extracted text, tables)
    """
    try:
        # Extract text
        text = extract_text_from_pdf_content(file_path, page_range)
        
        # Extract tables
        tables = extract_tables_from_pdf(file_path, page_range)
        
        return text, tables
        
    except Exception as e:
        logger.error(f"Error extracting from PDF: {str(e)}")
        # Fall back to Tika if extraction fails
        text = extract_text_with_tika(file_path, 'application/pdf')
        return text, []


def extract_text_from_pdf_content(file_path: str, page_range: Optional[Tuple[int, int]] = None) -> str:
    """
    Extract text content from a PDF file.
    
    Args:
        file_path: Path to the PDF file
        page_range: Optional tuple of (start_page, end_page) to extract (1-based)
        
    Returns:
        Extracted text
    """
    try:
        # Try to import PyPDF2
        import PyPDF2
        
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            
            # Determine page range
            if page_range:
                start_page, end_page = page_range
                # Convert to 0-based indexing
                start_idx = max(0, start_page - 1)
                end_idx = min(len(pdf_reader.pages), end_page)
                page_indices = range(start_idx, end_idx)
            else:
                page_indices = range(len(pdf_reader.pages))
            
            # Extract text from each page
            text_parts = []
            for i in page_indices:
                page = pdf_reader.pages[i]
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(f"--- Page {i+1} ---\n{page_text}")
            
            return "\n\n".join(text_parts)
        
    except ImportError:
        # Fall back to Tika if PyPDF2 is not available
        return extract_text_with_tika(file_path, 'application/pdf')


def extract_tables_from_pdf(file_path: str, page_range: Optional[Tuple[int, int]] = None) -> List[Dict[str, Any]]:
    """
    Extract tables from a PDF file.
    
    Args:
        file_path: Path to the PDF file
        page_range: Optional tuple of (start_page, end_page) to extract (1-based)
        
    Returns:
        List of extracted tables with metadata
    """
    tables = []
    
    # Try tabula-py first
    try:
        import tabula
        
        # Determine page range
        pages = None
        if page_range:
            start_page, end_page = page_range
            pages = ','.join(str(i) for i in range(start_page, end_page + 1))
        
        # Extract tables
        extracted_tables = tabula.read_pdf(
            file_path,
            pages=pages,
            multiple_tables=True,
            guess=True,
            lattice=True,
            stream=True
        )
        
        # Process extracted tables
        for i, df in enumerate(extracted_tables):
            if not df.empty:
                # Convert DataFrame to formatted text
                table_text = format_dataframe_as_text(df)
                
                # Get page number (approximate)
                page_num = i + 1 if not page_range else page_range[0] + i
                
                tables.append({
                    "content": table_text,
                    "page": page_num,
                    "rows": len(df),
                    "columns": len(df.columns),
                    "extraction_method": "tabula"
                })
        
    except Exception as e:
        logger.warning(f"Error extracting tables with tabula: {str(e)}")
        
        # Try camelot as fallback
        try:
            import camelot
            
            # Determine page range
            pages = "1-end"
            if page_range:
                start_page, end_page = page_range
                pages = f"{start_page}-{end_page}"
            
            # Extract tables
            extracted_tables = camelot.read_pdf(
                file_path,
                pages=pages,
                flavor='lattice'
            )
            
            # Process extracted tables
            for i, table in enumerate(extracted_tables):
                df = table.df
                if not df.empty:
                    # Convert DataFrame to formatted text
                    table_text = format_dataframe_as_text(df)
                    
                    tables.append({
                        "content": table_text,
                        "page": table.page,
                        "rows": len(df),
                        "columns": len(df.columns),
                        "extraction_method": "camelot"
                    })
                    
        except Exception as e2:
            logger.warning(f"Error extracting tables with camelot: {str(e2)}")
    
    return tables


def format_dataframe_as_text(df: pd.DataFrame) -> str:
    """
    Format a pandas DataFrame as a text table.
    
    Args:
        df: The DataFrame to format
        
    Returns:
        Formatted text representation of the table
    """
    # Clean column names
    df.columns = [str(col).strip() for col in df.columns]
    
    # Convert all values to strings
    for col in df.columns:
        df[col] = df[col].astype(str).apply(lambda x: x.strip())
    
    # Format as markdown table
    header = " | ".join(df.columns)
    separator = " | ".join(["---"] * len(df.columns))
    rows = [" | ".join(row) for _, row in df.iterrows()]
    
    return f"| {header} |\n| {separator} |\n" + "\n".join(f"| {row} |" for row in rows)


def extract_text_from_docx(file_path: str) -> Tuple[str, List[Dict[str, Any]]]:
    """
    Extract text and tables from a DOCX file.
    
    Args:
        file_path: Path to the DOCX file
        
    Returns:
        Tuple of (extracted text, tables)
    """
    try:
        # Try to import docx
        import docx
        
        doc = docx.Document(file_path)
        
        # Extract text from paragraphs
        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # Extract tables
        tables = []
        for i, table in enumerate(doc.tables):
            # Convert table to DataFrame
            data = []
            for j, row in enumerate(table.rows):
                data.append([cell.text for cell in row.cells])
            
            if data:
                # Create DataFrame
                df = pd.DataFrame(data[1:], columns=data[0] if data else None)
                
                # Format as text
                table_text = format_dataframe_as_text(df)
                
                tables.append({
                    "content": table_text,
                    "rows": len(data),
                    "columns": len(data[0]) if data and data[0] else 0,
                    "extraction_method": "python-docx"
                })
        
        return "\n\n".join(text_parts), tables
        
    except ImportError:
        # Fall back to Tika if docx is not available
        text = extract_text_with_tika(file_path, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        return text, []
    except Exception as e:
        logger.error(f"Error extracting from DOCX: {str(e)}")
        text = extract_text_with_tika(file_path, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        return text, []


def extract_text_from_txt(file_path: str) -> str:
    """
    Extract text from a plain text file.
    """
    with open(file_path, 'r', encoding='utf-8', errors='replace') as file:
        return file.read()


def extract_text_from_scan(file_path: str, blob_uri: str) -> Tuple[str, List[Dict[str, Any]]]:
    """
    Extract text and tables from a scanned document using Azure Form Recognizer.
    
    Args:
        file_path: Path to the scanned document
        blob_uri: Original blob URI
        
    Returns:
        Tuple of (extracted text, tables)
    """
    # Get Form Recognizer credentials from environment
    endpoint = os.environ.get("FORM_RECOGNIZER_ENDPOINT")
    key = os.environ.get("FORM_RECOGNIZER_KEY")
    
    if not endpoint or not key:
        logger.warning("Form Recognizer credentials not found, falling back to Tika")
        text = extract_text_with_tika(file_path, 'image/jpeg')
        return text, []
    
    try:
        # Import Azure Form Recognizer client
        from azure.ai.formrecognizer import DocumentAnalysisClient
        from azure.core.credentials import AzureKeyCredential
        
        # Create client
        document_analysis_client = DocumentAnalysisClient(
            endpoint=endpoint, 
            credential=AzureKeyCredential(key)
        )
        
        # Analyze document
        with open(file_path, "rb") as f:
            poller = document_analysis_client.begin_analyze_document(
                "prebuilt-document", document=f
            )
        result = poller.result()
        
        # Extract text from result
        text_parts = []
        for page in result.pages:
            page_text = []
            for line in page.lines:
                page_text.append(line.content)
            
            if page_text:
                text_parts.append(f"--- Page {page.page_number} ---\n" + "\n".join(page_text))
        
        # Extract tables
        tables = []
        for i, table in enumerate(result.tables):
            # Convert table to DataFrame
            data = []
            # Create empty grid
            rows = max(cell.row_index for cell in table.cells) + 1
            cols = max(cell.column_index for cell in table.cells) + 1
            grid = [['' for _ in range(cols)] for _ in range(rows)]
            
            # Fill in the grid
            for cell in table.cells:
                grid[cell.row_index][cell.column_index] = cell.content
            
            # Create DataFrame
            df = pd.DataFrame(grid[1:], columns=grid[0] if grid else None)
            
            # Format as text
            table_text = format_dataframe_as_text(df)
            
            tables.append({
                "content": table_text,
                "page": table.bounding_regions[0].page_number if table.bounding_regions else 1,
                "rows": rows,
                "columns": cols,
                "extraction_method": "form-recognizer"
            })
        
        return "\n\n".join(text_parts), tables
        
    except ImportError:
        logger.warning("Azure Form Recognizer SDK not found, falling back to Tika")
        text = extract_text_with_tika(file_path, 'image/jpeg')
        return text, []
    except Exception as e:
        logger.error(f"Error using Form Recognizer: {str(e)}")
        text = extract_text_with_tika(file_path, 'image/jpeg')
        return text, []


def extract_text_with_tika(file_path: str, content_type: str) -> str:
    """
    Extract text using Apache Tika REST API.
    """
    # Get Tika server URL from environment or use default
    tika_server_url = os.environ.get("TIKA_SERVER_URL", "http://tika:9998")
    
    try:
        with open(file_path, 'rb') as file:
            response = requests.put(
                f"{tika_server_url}/tika",
                headers={
                    'Accept': 'text/plain',
                    'Content-Type': content_type
                },
                data=file,
                timeout=60
            )
            
            if response.status_code == 200:
                return response.text
            else:
                raise Exception(f"Tika error: {response.status_code} - {response.text}")
    except Exception as e:
        logger.error(f"Error using Tika: {str(e)}")
        
        # If Tika fails, try a simple text extraction as last resort
        try:
            with open(file_path, 'r', encoding='utf-8', errors='replace') as file:
                return file.read()
        except:
            raise Exception(f"Failed to extract text with Tika and fallback methods: {str(e)}")


def combine_text_and_tables(text: str, tables: List[Dict[str, Any]]) -> str:
    """
    Combine extracted text and tables into a single document.
    
    Args:
        text: The extracted text
        tables: List of extracted tables
        
    Returns:
        Combined text with tables inserted at appropriate positions
    """
    if not tables:
        return text
    
    # Sort tables by page number
    tables_by_page = {}
    for table in tables:
        page = table.get("page", 1)
        if page not in tables_by_page:
            tables_by_page[page] = []
        tables_by_page[page].append(table)
    
    # Split text by page markers
    page_pattern = r'---\s*Page\s+(\d+)\s*---'
    page_matches = list(re.finditer(page_pattern, text))
    
    if not page_matches:
        # No page markers, just append tables at the end
        combined = text + "\n\n" + "--- EXTRACTED TABLES ---\n\n"
        for table in tables:
            combined += f"Table (Rows: {table.get('rows')}, Columns: {table.get('columns')}):\n"
            combined += table["content"] + "\n\n"
        return combined
    
    # Insert tables at the end of each page
    result_parts = []
    for i, match in enumerate(page_matches):
        page_num = int(match.group(1))
        
        # Get text for this page
        start_pos = match.end()
        end_pos = page_matches[i+1].start() if i+1 < len(page_matches) else len(text)
        page_text = text[start_pos:end_pos].strip()
        
        # Add page header
        result_parts.append(f"--- Page {page_num} ---")
        result_parts.append(page_text)
        
        # Add tables for this page
        if page_num in tables_by_page:
            result_parts.append(f"--- Tables on Page {page_num} ---")
            for table in tables_by_page[page_num]:
                result_parts.append(f"Table (Rows: {table.get('rows')}, Columns: {table.get('columns')}):")
                result_parts.append(table["content"])
    
    # Add any tables without page numbers at the end
    if 0 in tables_by_page or None in tables_by_page:
        result_parts.append("--- Additional Extracted Tables ---")
        for page_num in [0, None]:
            if page_num in tables_by_page:
                for table in tables_by_page[page_num]:
                    result_parts.append(f"Table (Rows: {table.get('rows')}, Columns: {table.get('columns')}):")
                    result_parts.append(table["content"])
    
    return "\n\n".join(result_parts)


def is_transient_error(error: Exception) -> bool:
    """
    Determine if an error is transient and should be retried.
    """
    # Check for common transient error patterns
    error_str = str(error).lower()
    
    # Network errors
    if any(msg in error_str for msg in ['timeout', 'connection', 'network', 'temporarily', 'throttled']):
        return True
    
    # Azure storage specific errors
    if any(msg in error_str for msg in ['server busy', 'operation could not be completed', '503', '500']):
        return True
    
    # Memory errors (common with large files)
    if any(msg in error_str for msg in ['memory', 'allocation', 'out of memory']):
        return True
    
    return False
